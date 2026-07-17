"""Output formatters: JSON (native), SRT, VTT, plain text, and Word (.docx)."""

from __future__ import annotations

import io
import json


def _ts_vtt(seconds: float) -> str:
    h, rem = divmod(max(0.0, seconds), 3600)
    m, s = divmod(rem, 60)
    return f"{int(h):02d}:{int(m):02d}:{s:06.3f}"


def _ts_srt(seconds: float) -> str:
    return _ts_vtt(seconds).replace(".", ",")


def to_json(result: dict) -> str:
    return json.dumps(result, ensure_ascii=False, indent=2)


def to_vtt(result: dict) -> str:
    lines = ["WEBVTT", ""]
    for seg in result["segments"]:
        lines.append(f"{_ts_vtt(seg['start'])} --> {_ts_vtt(seg['end'])}")
        lines.append(f"<v {seg['speaker']}>{seg['text']}")
        lines.append("")
    return "\n".join(lines)


def to_srt(result: dict) -> str:
    lines = []
    for i, seg in enumerate(result["segments"], start=1):
        lines.append(str(i))
        lines.append(f"{_ts_srt(seg['start'])} --> {_ts_srt(seg['end'])}")
        lines.append(f"[{seg['speaker']}] {seg['text']}")
        lines.append("")
    return "\n".join(lines)


def to_text(result: dict) -> str:
    return "\n".join(f"[{seg['speaker']}] {seg['text']}" for seg in result["segments"])


def to_docx(result: dict) -> bytes:
    """Word document with a 6-column table: ID, Start, End, Person, ':', Text."""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt

    # ID, Start, End, Person, ':', Transcript — the ':' column is a separator,
    # so keep it hairline; the transcript takes the remaining width.
    _DOCX_COL_WIDTHS = {
        0: Cm(1.0), 1: Cm(2.2), 2: Cm(2.2), 3: Cm(3.0), 4: Cm(0.5), 5: Cm(8.1),
    }

    doc = Document()
    doc.add_heading("Transcript", level=1)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Duration: {_ts_vtt(result.get('duration', 0))}   ·   "
        f"Language: {result.get('language', 'auto')}   ·   "
        f"Speakers: {', '.join(result.get('speakers', []))}"
    ).font.size = Pt(9)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0].cells
    for i, title in enumerate(["ID", "Start", "End", "Person", "", "Transcript"]):
        run = header[i].paragraphs[0].add_run(title)
        run.bold = True

    for i, seg in enumerate(result["segments"], start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = _ts_vtt(seg["start"])
        row[2].text = _ts_vtt(seg["end"])
        row[3].text = seg["speaker"]
        row[4].text = ":"
        row[5].text = seg["text"]

    # Word only honours column widths when autofit is off and the width is set
    # on every cell of the column, not just on the column object.
    table.autofit = False
    for idx, width in _DOCX_COL_WIDTHS.items():
        table.columns[idx].width = width
        for row in table.rows:
            row.cells[idx].width = width

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


FORMATTERS = {
    "json": (to_json, "application/json"),
    "verbose_json": (to_json, "application/json"),
    "vtt": (to_vtt, "text/vtt"),
    "srt": (to_srt, "application/x-subrip"),
    "text": (to_text, "text/plain"),
    "docx": (to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
}
